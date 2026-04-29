"""一次讀 3 份 docx 報告,輸出純文字到指定資料夾,供下一輪閱讀。"""
import sys
try:
    sys.stdout.reconfigure(encoding="utf-8")  # type: ignore[attr-defined]
except Exception:
    pass
from pathlib import Path
from docx import Document

SRC_DIR = Path(r"D:\下載onedrive\drive-download-20260429T101504Z-3-001")
OUT_DIR = Path(__file__).resolve().parent / "_external_reports"
OUT_DIR.mkdir(exist_ok=True)

FILES = [
    "gov-expense-kb_disclaimer-design-report.md.docx",
    "gov-expense-kb_scenario-automation-report.md.docx",
    "gov-expense-kb_tag-linking-recommendations.md.docx",
]


def docx_to_text(path: Path) -> str:
    doc = Document(str(path))
    out: list[str] = []
    for p in doc.paragraphs:
        style = (p.style.name if p.style else "") or "Normal"
        text = p.text
        if not text.strip():
            out.append("")
            continue
        # 對 heading 加 markdown #,讓結構保留
        if style.startswith("Heading"):
            try:
                level = int(style.replace("Heading", "").strip() or "1")
            except ValueError:
                level = 1
            out.append("#" * max(1, min(level, 6)) + " " + text.strip())
        elif style == "Title":
            out.append("# " + text.strip())
        else:
            out.append(text)
    # 表格
    for i, t in enumerate(doc.tables):
        out.append(f"\n[Table {i+1}]")
        for row in t.rows:
            cells = [c.text.replace("\n", " / ").strip() for c in row.cells]
            out.append(" | ".join(cells))
    return "\n".join(out)


for fname in FILES:
    src = SRC_DIR / fname
    if not src.exists():
        print(f"⚠ 找不到 {src}")
        continue
    text = docx_to_text(src)
    out_name = fname.replace(".md.docx", ".txt")
    (OUT_DIR / out_name).write_text(text, encoding="utf-8", newline="\n")
    print(f"✓ {fname} → {out_name} ({len(text)} chars)")

print(f"\n輸出資料夾:{OUT_DIR}")
