"""01_extract.py — 來源檔抽取為純文字。

職責
----
把 ``00_source/{機關}/`` 下的 PDF / DOCX / MD 轉成乾淨純文字,
依 (類別, 母題) 分類存放到 ``01_extracted/{類別}/{母題}/{slug}.txt``,
並輸出同名 ``.meta.json`` 紀錄抽取資訊。

非職責(屬於 02_parse.py / 03_build_index.py)
----------------------------------------------
- 條文切分、Q&A 切題
- 產生 MD front-matter
- 建立 JSON 索引

使用範例
--------
    # 全跑(冪等,輸出較新會跳過)
    python 05_scripts/01_extract.py

    # 只跑單檔
    python 05_scripts/01_extract.py --file 國內出差旅費報支要點_第五條_*.md

    # 強制覆寫
    python 05_scripts/01_extract.py --force

    # 掃描型 PDF 走 OCR(首次會下載模型,慢)
    python 05_scripts/01_extract.py --file scanned.pdf --ocr

    # 預演,不實際輸出
    python 05_scripts/01_extract.py --dry-run

退出代碼
--------
0  全部成功
1  啟動環境錯誤(找不到 source 目錄等)
2  有檔案處理失敗
"""

from __future__ import annotations

import argparse
import csv
import json
import logging
import re
import sys
from dataclasses import asdict, dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Literal, Optional

# ─────────────────────────────────────────────
# 路徑與常數
# ─────────────────────────────────────────────

ROOT = Path(__file__).resolve().parent.parent
SOURCE_DIR = ROOT / "00_source"
EXTRACT_DIR = ROOT / "01_extracted"
MANIFEST_PATH = SOURCE_DIR / "_manifest.csv"
SKIP_LIST_PATH = SOURCE_DIR / "_skip.txt"

CATEGORY_DIR_NAMES: dict[str, str] = {
    "A": "A_核心法規",
    "B": "B_支出標準",
    "C": "C_解釋函令",
    "D": "D_問答集",
}
UNSORTED = "_unsorted"

# 文件類型 → 類別代碼(順序重要,先比對特殊類型)
DOC_TYPE_PATTERNS: list[tuple[re.Pattern[str], str]] = [
    (re.compile(r"問答集|Q\s*&\s*A|QA"), "D"),
    (re.compile(r"函釋|解釋(?:令|函|彙編)|釋示"), "C"),
    (re.compile(r"(?:標準|額度|數額|費率)表"), "B"),
    (re.compile(r"要點|辦法|規則|原則|注意事項|處理|綱要|基準"), "A"),
]

# 母題關鍵字對照(順序重要,長詞優先;比對來源欄位 + 檔名)
# 注意:
# 1) 「支出憑證」「鐘點費」「出席費」等費目名稱常被其他法規條文/問答內文
#    提及,單以這類詞做關鍵字會誤判;故 PARENT_KEYWORDS 只列「特定法規/
#    問答集名稱」的識別字串,不列泛費目詞。
# 2) 順序由最具體(法規/問答集前綴)到較通用(母題本名)。
PARENT_KEYWORDS: list[tuple[str, str]] = [
    ("經費結報常見疑義問答集", "支出憑證與結報"),
    ("經費結報", "支出憑證與結報"),
    ("政府支出憑證處理要點", "支出憑證與結報"),
    ("會計憑證", "支出憑證與結報"),
    ("國內出差", "國內旅費"),
    ("國內旅費", "國內旅費"),
    ("國外出差", "國外旅費"),
    ("國外旅費", "國外旅費"),
    ("派赴國外", "國外旅費"),
    ("派赴大陸", "國外旅費"),
    ("講座鐘點費", "講座鐘點費"),
    ("鐘點費", "講座鐘點費"),
    ("出席費", "酬勞費"),
    ("稿費", "酬勞費"),
    ("審查費", "酬勞費"),
    ("國外顧問", "國外專家"),
    ("國外專家", "國外專家"),
    ("聘請國外", "國外專家"),
    ("國科會", "國科會專章"),
    ("國家科學及技術委員會", "國科會專章"),
    ("專題研究計畫", "國科會專章"),
    ("教育部", "教育部專章"),
    ("校務基金", "教育部專章"),
    ("產學合作", "教育部專章"),
    ("行動電話", "其他"),
    ("會議費", "其他"),
    ("訓練", "其他"),
    ("講習", "其他"),
]

# 觸發 OCR 建議的閾值(平均每頁字數)
MIN_CHARS_PER_PAGE = 50

SUPPORTED_SUFFIXES = {".pdf", ".docx", ".md", ".markdown", ".txt"}


# ─────────────────────────────────────────────
# 結果結構
# ─────────────────────────────────────────────

ExtractMethod = Literal[
    "pdfplumber",
    "ocr",
    "docx",
    "markdown_passthrough",
    "skipped",
    "failed",
    "unsupported",
]


@dataclass
class ExtractResult:
    """單檔抽取結果。"""

    source_path: str
    output_path: Optional[str] = None
    method: ExtractMethod = "failed"
    char_count: int = 0
    page_count: int = 0
    category: Optional[str] = None
    parent: Optional[str] = None
    warnings: list[str] = field(default_factory=list)
    error: Optional[str] = None
    extracted_at: str = ""
    source_metadata: dict[str, str] = field(default_factory=dict)

    @property
    def ok(self) -> bool:
        return self.error is None and self.method not in {"failed", "unsupported"}


# ─────────────────────────────────────────────
# Manifest
# ─────────────────────────────────────────────


def load_manifest(path: Path) -> dict[str, dict[str, str]]:
    """讀取 _manifest.csv,以 filename 為 key。"""
    if not path.exists():
        return {}
    rows: dict[str, dict[str, str]] = {}
    with path.open(encoding="utf-8-sig", newline="") as f:
        reader = csv.DictReader(f)
        for row in reader:
            fname = (row.get("filename") or "").strip()
            if not fname:
                continue
            rows[fname] = {k: (v or "").strip() for k, v in row.items()}
    return rows


# ─────────────────────────────────────────────
# 路由判斷(類別 / 母題)
# ─────────────────────────────────────────────


def infer_category(doc_type_text: str) -> Optional[str]:
    """從文件類型字串(如「要點」、「問答集_QA」)推斷類別代碼。"""
    if not doc_type_text:
        return None
    for pattern, code in DOC_TYPE_PATTERNS:
        if pattern.search(doc_type_text):
            return code
    return None


def infer_parent(*texts: str) -> Optional[str]:
    """從多段文字(來源欄位、檔名等)推斷母題。"""
    blob = " ".join(t for t in texts if t)
    for keyword, parent in PARENT_KEYWORDS:
        if keyword in blob:
            return parent
    return None


def determine_routing(
    src: Path,
    md_metadata: dict[str, str],
    manifest_row: Optional[dict[str, str]],
) -> tuple[Optional[str], Optional[str], list[str]]:
    """決定 (category_code, parent, warnings)。manifest 優先,MD metadata 次之。"""
    warnings: list[str] = []
    category: Optional[str] = None
    parent: Optional[str] = None

    if manifest_row:
        category = manifest_row.get("category") or None
        parent = manifest_row.get("parent") or None

    if not category:
        doc_type = md_metadata.get("文件類型", "")
        category = infer_category(doc_type)
        if category:
            warnings.append(f"類別由 MD 文件類型欄位推斷:{doc_type} → {category}")

    if not parent:
        parent = infer_parent(
            md_metadata.get("來源法規", ""),
            md_metadata.get("__h1__", ""),
            src.name,
        )
        if parent:
            warnings.append(f"母題由內容/檔名推斷:{parent}")

    return category, parent, warnings


def output_path_for(category: Optional[str], parent: Optional[str], slug: str) -> Path:
    """組出輸出路徑;類別或母題缺失時放入 _unsorted/。"""
    cat_dir = CATEGORY_DIR_NAMES.get(category, UNSORTED) if category else UNSORTED
    parent_dir = parent or UNSORTED
    return EXTRACT_DIR / cat_dir / parent_dir / f"{slug}.txt"


# ─────────────────────────────────────────────
# 抽取器(各格式)
# ─────────────────────────────────────────────


def extract_pdf(path: Path, log: logging.Logger) -> tuple[str, int]:
    """用 pdfplumber 抽 PDF 文字,回傳 (text, page_count)。"""
    try:
        import pdfplumber
    except ImportError as e:
        raise RuntimeError("缺少 pdfplumber,請執行 pip install pdfplumber") from e

    pages_text: list[str] = []
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages, 1):
            try:
                pages_text.append(page.extract_text() or "")
            except Exception as e:  # 個別頁失敗不擋整檔
                log.warning(f"  第 {i} 頁抽取失敗:{e}")
                pages_text.append("")
    return "\n\n".join(pages_text), len(pages_text)


def extract_pdf_ocr(path: Path, log: logging.Logger) -> tuple[str, int]:
    """用 paddleocr + pdf2image OCR 掃描型 PDF。"""
    try:
        from paddleocr import PaddleOCR
    except ImportError as e:
        raise RuntimeError(
            "OCR 需要 paddleocr,請執行 pip install paddleocr paddlepaddle"
        ) from e
    try:
        from pdf2image import convert_from_path
    except ImportError as e:
        raise RuntimeError("OCR 需要 pdf2image,請執行 pip install pdf2image") from e

    log.info("  初始化 PaddleOCR(首次會下載模型,可能需數分鐘)")
    ocr = PaddleOCR(use_angle_cls=True, lang="ch", show_log=False)
    images = convert_from_path(str(path), dpi=200)
    pages_text: list[str] = []
    for i, img in enumerate(images, 1):
        log.info(f"  OCR 第 {i}/{len(images)} 頁")
        result = ocr.ocr(img, cls=True)
        if not result or not result[0]:
            pages_text.append("")
            continue
        lines = [line[1][0] for line in result[0] if line and line[1]]
        pages_text.append("\n".join(lines))
    return "\n\n".join(pages_text), len(pages_text)


def extract_docx(path: Path) -> tuple[str, int]:
    """用 python-docx 抽段落與表格,回傳 (text, paragraph_count)。"""
    try:
        from docx import Document
    except ImportError as e:
        raise RuntimeError("缺少 python-docx,請執行 pip install python-docx") from e

    doc = Document(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            line = " | ".join(c for c in cells if c)
            if line:
                parts.append(line)
    return "\n".join(parts), len(doc.paragraphs)


def extract_md(path: Path) -> tuple[str, dict[str, str]]:
    """直通 MD,剝掉檔頭 metadata 區塊,回傳 (正文, metadata)。

    來源 MD 開頭格式(範例):
        # XXX標題
        **來源法規:** ...
        **文件類型:** ...
        **條次:** ...
        **標籤:** #tag1 #tag2
        ---
        (正文)
    """
    text = path.read_text(encoding="utf-8")
    return strip_md_header(text)


def strip_md_header(text: str) -> tuple[str, dict[str, str]]:
    """剝離 MD 開頭的標題與 ``**欄位:**`` 區塊,回傳 (正文, metadata)。"""
    metadata: dict[str, str] = {}
    field_pat = re.compile(r"^\*\*\s*([^*：:]+?)\s*[：:]\s*\*\*\s*(.+?)\s*$")
    hr_pat = re.compile(r"^[-*_]{3,}\s*$")

    # 來源 MD 由 PDF 抽取轉成,常夾雜 \f(form feed)當頁碼分隔。
    # \f 會被 splitlines 視為分行,造成 H1 與 metadata 區塊被截斷。
    # 直接移除以還原邏輯行。
    text = text.replace("\f", " ")

    lines = text.splitlines()
    i = 0
    in_header = True

    # 頭區塊:消化開頭的空白行 → H1 → metadata 欄位 → 第一條 ---
    while i < len(lines) and in_header:
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # 第一個 H1 視為標題,不計入正文
        if "__h1__" not in metadata and stripped.startswith("# "):
            metadata["__h1__"] = stripped[2:].strip()
            i += 1
            continue

        m = field_pat.match(stripped)
        if m:
            key = m.group(1).strip()
            value = m.group(2).strip()
            # 標籤欄位特別處理:把 "#a #b" 攤平成 "a, b"
            if key == "標籤":
                tags = re.findall(r"#([^\s#]+)", value)
                metadata[key] = ", ".join(tags) if tags else value
            else:
                metadata[key] = value
            i += 1
            continue

        # 遇到分隔線就結束 header,並消化掉它
        if hr_pat.match(stripped):
            i += 1
            in_header = False
            break

        # 既非 H1、欄位、分隔線、空行 → 已進入正文
        in_header = False
        break

    body_lines = lines[i:]
    # 移除尾端空白行
    while body_lines and not body_lines[-1].strip():
        body_lines.pop()
    # 移除前段多餘空白行
    while body_lines and not body_lines[0].strip():
        body_lines.pop(0)

    body = "\n".join(body_lines)
    if body and not body.endswith("\n"):
        body += "\n"
    return body, metadata


# ─────────────────────────────────────────────
# 文字後處理
# ─────────────────────────────────────────────


def reflow_pdf_text(text: str) -> str:
    """PDF 抽取常因頁寬硬斷行,把不該斷的中文行併接。

    規則:若上一行最後一字是中文且非句末標點,且本行起始也是中文,
    則合併兩行。保留段落間空白行。
    """
    sentence_end = set("。!?!?;;…」』")
    out: list[str] = []
    for ln in text.splitlines():
        if not out or not ln.strip() or not out[-1].strip():
            out.append(ln)
            continue
        prev = out[-1]
        last_char = prev.rstrip()[-1] if prev.rstrip() else ""
        first_char = ln.lstrip()[0] if ln.lstrip() else ""
        if (
            _is_cjk(last_char)
            and last_char not in sentence_end
            and _is_cjk(first_char)
        ):
            out[-1] = prev.rstrip() + ln.lstrip()
        else:
            out.append(ln)
    return "\n".join(out)


def _is_cjk(ch: str) -> bool:
    """檢查單字元是否為 CJK。"""
    if not ch:
        return False
    cp = ord(ch)
    return (
        0x4E00 <= cp <= 0x9FFF
        or 0x3400 <= cp <= 0x4DBF
        or 0xF900 <= cp <= 0xFAFF
    )


# ─────────────────────────────────────────────
# 檔名工具
# ─────────────────────────────────────────────


def slugify(filename: str) -> str:
    """檔名轉安全 slug。保留中文,移除路徑非法字元。"""
    base = Path(filename).stem
    base = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", base)
    base = base.strip(" ._")
    return base or "untitled"


# ─────────────────────────────────────────────
# 主流程
# ─────────────────────────────────────────────


def should_skip(src: Path, dst: Path, force: bool) -> bool:
    """輸出較新且未 --force 時跳過。"""
    if force or not dst.exists():
        return False
    return dst.stat().st_mtime >= src.stat().st_mtime


def process_file(
    src: Path,
    manifest: dict[str, dict[str, str]],
    *,
    use_ocr: bool,
    force: bool,
    dry_run: bool,
    log: logging.Logger,
) -> ExtractResult:
    """處理單一來源檔。"""
    rel_src = src.relative_to(ROOT).as_posix()
    result = ExtractResult(
        source_path=rel_src,
        extracted_at=datetime.now().isoformat(timespec="seconds"),
    )

    suffix = src.suffix.lower()
    if suffix not in SUPPORTED_SUFFIXES:
        result.method = "unsupported"
        result.error = f"不支援的副檔名:{suffix}"
        log.warning(f"跳過 {src.name}:{result.error}")
        return result

    manifest_row = manifest.get(src.name)
    if not manifest_row:
        result.warnings.append("未列於 _manifest.csv,將靠檔名/內容推斷類別與母題")

    try:
        # ── 抽取文字 ──
        md_metadata: dict[str, str] = {}
        if suffix == ".pdf":
            if use_ocr:
                text, page_count = extract_pdf_ocr(src, log)
                result.method = "ocr"
            else:
                text, page_count = extract_pdf(src, log)
                result.method = "pdfplumber"
            text = reflow_pdf_text(text)
            result.page_count = page_count
        elif suffix == ".docx":
            text, page_count = extract_docx(src)
            result.method = "docx"
            result.page_count = page_count
        else:  # .md / .markdown / .txt
            text, md_metadata = extract_md(src)
            result.method = "markdown_passthrough"
            result.page_count = 1
            result.source_metadata = md_metadata

        result.char_count = len(text)

        # ── 路由判斷 ──
        category, parent, route_warnings = determine_routing(src, md_metadata, manifest_row)
        result.category = category
        result.parent = parent
        result.warnings.extend(route_warnings)
        if not category:
            result.warnings.append("未能判斷類別,放入 _unsorted/(可於 _manifest.csv 補上)")
        if not parent:
            result.warnings.append("未能判斷母題,放入 _unsorted/(可於 _manifest.csv 補上)")

        # ── OCR 建議 ──
        if suffix == ".pdf" and not use_ocr and result.page_count > 0:
            avg = result.char_count / result.page_count
            if avg < MIN_CHARS_PER_PAGE:
                result.warnings.append(
                    f"平均每頁僅 {avg:.0f} 字,可能為掃描型 PDF,建議加 --ocr 重跑"
                )

        # ── 截斷標題警告(觀察:來源 MD 標題常被檔名截短) ──
        h1 = md_metadata.get("__h1__", "")
        if h1 and not re.search(r"[。?!?!」』]\s*$", h1) and len(h1) >= 28:
            result.warnings.append(f"H1 標題疑似被截斷,後續切分階段請從本文重抽:{h1}")

        # ── 寫檔 ──
        slug = slugify(src.name)
        dst = output_path_for(category, parent, slug)

        if should_skip(src, dst, force):
            result.method = "skipped"
            result.output_path = dst.relative_to(ROOT).as_posix()
            log.info(f"跳過 {src.name}(輸出較新,加 --force 可覆寫)")
            return result

        result.output_path = dst.relative_to(ROOT).as_posix()

        if not dry_run:
            dst.parent.mkdir(parents=True, exist_ok=True)
            dst.write_text(text, encoding="utf-8", newline="\n")
            meta_dict = asdict(result)
            meta_dict["manifest_row"] = manifest_row
            dst.with_suffix(".meta.json").write_text(
                json.dumps(meta_dict, ensure_ascii=False, indent=2),
                encoding="utf-8",
                newline="\n",
            )

        log.info(
            f"{'[dry-run] ' if dry_run else ''}完成 {src.name} → "
            f"{result.output_path}({result.method},{result.char_count} 字)"
        )

    except Exception as e:
        result.error = str(e)
        result.method = "failed"
        log.error(f"處理失敗 {src.name}:{e}")

    return result


def load_skip_set(path: Path) -> set[str]:
    """讀 _skip.txt(每行一個檔名);註解行(#) 與空行忽略。"""
    if not path.exists():
        return set()
    out: set[str] = set()
    for line in path.read_text(encoding="utf-8").splitlines():
        line = line.strip()
        if not line or line.startswith("#"):
            continue
        out.add(line)
    return out


def find_inputs(filter_arg: Optional[str]) -> list[Path]:
    """蒐集要處理的來源檔。filter_arg 可為檔名、glob 或路徑。

    自動跳過 _skip.txt 中列出的檔名(常用於排除重複/雜訊版本)。
    --file 明確指定檔案時,不套用 skip 過濾(讓人工 override)。
    """
    if filter_arg:
        # 絕對路徑
        p = Path(filter_arg)
        if p.is_absolute() and p.exists() and p.is_file():
            return [p]
        # glob 比對(支援 *)
        matches = list(SOURCE_DIR.rglob(filter_arg))
        # 若沒有 wildcard 又找不到,試純檔名
        if not matches and "*" not in filter_arg:
            matches = list(SOURCE_DIR.rglob(p.name))
        files = [m for m in matches if m.is_file() and m.suffix.lower() in SUPPORTED_SUFFIXES]
        if not files:
            raise FileNotFoundError(f"在 00_source/ 找不到符合 {filter_arg!r} 的檔案")
        return sorted(files)

    skip = load_skip_set(SKIP_LIST_PATH)
    return sorted(
        p
        for p in SOURCE_DIR.rglob("*")
        if p.is_file()
        and p.suffix.lower() in SUPPORTED_SUFFIXES
        and not p.name.startswith("_")
        and p.name not in skip
    )


def print_summary(results: list[ExtractResult], dry_run: bool) -> None:
    """印彙整表與警告清單。"""
    if not results:
        return
    print()
    print("=" * 78)
    header = f"{'檔名':<42}{'方法':<22}{'類別':<6}{'母題':<10}{'字數':>6}"
    print(header)
    print("-" * 78)

    counts: dict[str, int] = {}
    for r in results:
        name = Path(r.source_path).name
        if _display_width(name) > 40:
            name = _truncate_display(name, 40)
        category = r.category or "-"
        parent = r.parent or "-"
        print(
            f"{_pad_display(name, 42)}"
            f"{r.method:<22}"
            f"{category:<6}"
            f"{_pad_display(parent, 10)}"
            f"{r.char_count:>6}"
        )
        counts[r.method] = counts.get(r.method, 0) + 1

    print("-" * 78)
    summary = ", ".join(f"{k}={v}" for k, v in sorted(counts.items()))
    tag = " (dry-run)" if dry_run else ""
    print(f"總計 {len(results)} 檔{tag}: {summary}")

    warned = [r for r in results if r.warnings]
    if warned:
        print(f"\n有警告的檔案({len(warned)}):")
        for r in warned:
            print(f"  • {Path(r.source_path).name}")
            for w in r.warnings:
                print(f"      - {w}")

    failed = [r for r in results if not r.ok]
    if failed:
        print(f"\n失敗 / 未處理({len(failed)}):")
        for r in failed:
            print(f"  • {Path(r.source_path).name}: {r.error or r.method}")


def _display_width(s: str) -> int:
    """估計顯示寬度:CJK 視為 2,其餘 1。"""
    return sum(2 if _is_cjk(c) or 0x3000 <= ord(c) <= 0x303F else 1 for c in s)


def _pad_display(s: str, width: int) -> str:
    """依顯示寬度補空格。"""
    pad = max(0, width - _display_width(s))
    return s + " " * pad


def _truncate_display(s: str, width: int) -> str:
    """依顯示寬度截斷,結尾加 …。"""
    out: list[str] = []
    used = 0
    for ch in s:
        w = 2 if _is_cjk(ch) or 0x3000 <= ord(ch) <= 0x303F else 1
        if used + w > width - 1:
            break
        out.append(ch)
        used += w
    return "".join(out) + "…"


def main(argv: Optional[list[str]] = None) -> int:
    parser = argparse.ArgumentParser(
        description="把 00_source/ 內的 PDF/DOCX/MD 抽取為純文字到 01_extracted/",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    parser.add_argument(
        "--file",
        help="只處理指定檔(檔名、glob 或絕對路徑);省略則跑全部",
    )
    parser.add_argument("--ocr", action="store_true", help="PDF 走 OCR(慢、需要模型)")
    parser.add_argument("--force", action="store_true", help="即使輸出較新仍覆寫")
    parser.add_argument("--dry-run", action="store_true", help="只列出將要做什麼,不寫檔")
    parser.add_argument("-v", "--verbose", action="store_true", help="顯示 DEBUG 日誌")
    args = parser.parse_args(argv)

    logging.basicConfig(
        level=logging.DEBUG if args.verbose else logging.INFO,
        format="%(levelname)s %(message)s",
    )
    log = logging.getLogger("extract")

    if not SOURCE_DIR.exists():
        log.error(f"找不到來源目錄:{SOURCE_DIR}")
        return 1

    EXTRACT_DIR.mkdir(parents=True, exist_ok=True)

    manifest = load_manifest(MANIFEST_PATH)
    if not manifest:
        log.warning(f"找不到或無法讀取 {MANIFEST_PATH.name},僅靠檔名/內容推斷")

    try:
        inputs = find_inputs(args.file)
    except FileNotFoundError as e:
        log.error(str(e))
        return 1

    if not inputs:
        log.warning("沒有可處理的檔案")
        return 0

    log.info(f"準備處理 {len(inputs)} 個檔案" + (" (dry-run)" if args.dry_run else ""))

    results: list[ExtractResult] = []
    for src in inputs:
        results.append(
            process_file(
                src,
                manifest,
                use_ocr=args.ocr,
                force=args.force,
                dry_run=args.dry_run,
                log=log,
            )
        )

    print_summary(results, args.dry_run)

    n_failed = sum(1 for r in results if not r.ok)
    return 0 if n_failed == 0 else 2


if __name__ == "__main__":
    sys.exit(main())
