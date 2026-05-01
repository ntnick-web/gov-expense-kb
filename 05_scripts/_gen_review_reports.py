"""一次性腳本:產出兩份檢視報告(2026-05-01)
- 報告 1:資料庫處理流程優化報告
- 報告 2:資料庫資料夾整理報告

執行:python 05_scripts/_gen_review_reports.py
輸出:docs/_資料庫處理流程優化報告_2026-05-01.docx
      docs/_資料庫資料夾整理報告_2026-05-01.docx
"""
from __future__ import annotations
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
OUT_DIR = ROOT / "docs"

# ─────────────────────────────────────────────────────────────
# 共用樣式 helpers
# ─────────────────────────────────────────────────────────────

FONT_HAN = "Microsoft JhengHei"  # 微軟正黑體
FONT_EN = "Calibri"


def _set_cell_bg(cell, color_hex: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def _set_run_font(run, size_pt=11, bold=False, color=None):
    run.font.name = FONT_EN
    rpr = run._element.rPr
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        run._element.insert(0, rpr)
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), FONT_HAN)
    rfonts.set(qn("w:ascii"), FONT_EN)
    rfonts.set(qn("w:hAnsi"), FONT_EN)
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def doc_init() -> Document:
    d = Document()
    # base style
    style = d.styles["Normal"]
    style.font.name = FONT_EN
    style.font.size = Pt(10.5)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), FONT_HAN)
    rfonts.set(qn("w:ascii"), FONT_EN)
    rfonts.set(qn("w:hAnsi"), FONT_EN)

    # margins
    for section in d.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
    return d


def add_heading(d: Document, text: str, level: int):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(4)
    if level == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        _set_run_font(run, size_pt=22, bold=True, color="2F4858")
    elif level == 1:
        run = p.add_run(text)
        _set_run_font(run, size_pt=15, bold=True, color="2F4858")
        # bottom border
        pPr = p._element.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "8")
        bottom.set(qn("w:color"), "8FA8B8")
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == 2:
        run = p.add_run(text)
        _set_run_font(run, size_pt=13, bold=True, color="3B6E8F")
    elif level == 3:
        run = p.add_run(text)
        _set_run_font(run, size_pt=11.5, bold=True, color="555555")
    return p


def add_para(d: Document, text: str, bold=False, italic=False, size=10.5, color=None,
             align=None, indent_cm=None):
    p = d.add_paragraph()
    if align:
        p.alignment = align
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    _set_run_font(run, size_pt=size, bold=bold, color=color)
    run.italic = italic
    return p


def add_bullet(d: Document, text: str, level: int = 0, color=None):
    p = d.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.6)
    p.paragraph_format.space_after = Pt(2)
    # need to re-add as the style wipes; simpler: just use a plain para with •
    p.text = ""
    bullet_char = ["•", "◦", "▪"][min(level, 2)]
    run = p.add_run(f"{bullet_char}  {text}")
    _set_run_font(run, size_pt=10.5, color=color)
    return p


def add_numbered(d: Document, text: str, n: int, color=None):
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"{n}. {text}")
    _set_run_font(run, size_pt=10.5, color=color)
    return p


def add_callout(d: Document, title: str, body: str, bg="FFF7E6", border="E8B85E"):
    """Single-cell shaded callout box."""
    table = d.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    _set_cell_bg(cell, bg)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    # set borders
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "6")
        b.set(qn("w:color"), border)
        tcBorders.append(b)
    tcPr.append(tcBorders)

    # title
    p_title = cell.paragraphs[0]
    p_title.paragraph_format.space_after = Pt(2)
    r = p_title.add_run(title)
    _set_run_font(r, size_pt=11, bold=True, color="6B4A1F")
    # body
    p_body = cell.add_paragraph()
    p_body.paragraph_format.space_after = Pt(2)
    r2 = p_body.add_run(body)
    _set_run_font(r2, size_pt=10, color="3D2E13")
    d.add_paragraph()  # spacer


def add_table(d: Document, header: list[str], rows: list[list[str]],
              col_widths_cm: list[float] | None = None,
              header_bg="3B6E8F", header_fg="FFFFFF", zebra=True):
    table = d.add_table(rows=1 + len(rows), cols=len(header))
    table.autofit = False
    table.style = "Light Grid Accent 1"

    if col_widths_cm:
        for col_idx, w in enumerate(col_widths_cm):
            for row in table.rows:
                row.cells[col_idx].width = Cm(w)

    # header
    for i, h in enumerate(header):
        cell = table.cell(0, i)
        _set_cell_bg(cell, header_bg)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        _set_run_font(r, size_pt=10, bold=True, color=header_fg)

    # body
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.cell(ri + 1, ci)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            if zebra and ri % 2 == 1:
                _set_cell_bg(cell, "F4F7FA")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(val))
            _set_run_font(r, size_pt=9.5)
    return table


def add_ascii_block(d: Document, text: str):
    """Monospace ASCII block(流程圖 / 心智圖)"""
    table = d.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    _set_cell_bg(cell, "F8FAFC")
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "C9D5DE")
        tcBorders.append(b)
    tcPr.append(tcBorders)

    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for i, line in enumerate(text.splitlines()):
        if i > 0:
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        rpr = run._element.rPr
        if rpr is None:
            rpr = OxmlElement("w:rPr")
            run._element.insert(0, rpr)
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:eastAsia"), "Consolas")
        rfonts.set(qn("w:ascii"), "Consolas")
        rfonts.set(qn("w:hAnsi"), "Consolas")
    d.add_paragraph()


# ─────────────────────────────────────────────────────────────
# 報告 1:資料庫處理流程優化報告
# ─────────────────────────────────────────────────────────────

def build_flow_report(out_path: Path):
    d = doc_init()

    # 封面
    add_heading(d, "資料庫處理流程優化報告", 0)
    add_para(d, "── 政府支出法規知識庫 ──", size=12, color="6A7B8C",
             align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(d, "資料蒐集 → 抽取 → 解構 → 整理 → 索引 → 視覺化",
             size=10.5, color="6A7B8C", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(d, "完整處理鏈、決策點、控管條件、核心邏輯與護城河分析",
             size=10.5, italic=True, color="6A7B8C", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(d, "報告日期:2026-05-01    DATA_VERSION:2026-05-01a    節點數:520    情境卡:73 可見",
             size=9.5, color="888888", align=WD_ALIGN_PARAGRAPH.CENTER)
    d.add_paragraph()

    # 執行摘要
    add_heading(d, "執行摘要", 1)
    add_para(d, "本資料庫採「單向資料流 + SSOT(02_markdown)+ 純靜態前端」三層架構,從政府公開法規 PDF 抽取 → 解構 → 結構化 MD → JSON 索引 → 純前端視覺化,無資料庫、無框架、無 CDN 依賴。整體流程歷經 8 個明確階段,每階段皆有獨立工具、控管條件與決策點。")
    add_para(d, "目前涵蓋 3 個母題(國內旅費、國外旅費、支出憑證與結報),520 節點、73 張可見情境卡、6 個情境樹 root、482 條人工/推斷雙模式關聯邊;前端含 6 視圖(Landing 三入口、情境檢索、條文庫、試算表、抽屜、比較模式)、馬卡龍 7 色 token 設計系統、城市 fallback、保險表互動 widget、信度三層免責 UI 等獨家功能。")
    add_para(d, "本報告將上述架構拆解為「8 階段流水線」並標出 27 個關鍵決策點與 18 條控管條件,接續分析核心邏輯與護城河,並提出短中長期改善建議共 28 項。", bold=True)

    # 一、心智圖
    add_heading(d, "一、整體處理流程心智圖", 1)
    add_para(d, "下圖為從原始檔到上線網站的完整資料流。實線為主管線,虛線為輔助/校對工具,雙線為決策控管點(◆)。")
    add_ascii_block(d, """\
              ┌──────────────────────────────────────────────────┐
              │   ◆ 中立角色原則(法源位階 A>B>C>D,不給判斷)   │
              │   ◆ SSOT = 02_markdown(其他皆可重建)            │
              │   ◆ 純靜態無框架(GitHub Pages + CF Analytics)   │
              └──────────────────────────────────────────────────┘
                                     │ 全域控管原則
                                     ▼
  ┌────────┐    ┌────────┐    ┌────────┐    ┌────────┐    ┌────────┐
  │00_source│───▶│01_extr.│───▶│02_md   │───▶│03_index│───▶│04_web  │
  │PDF/DOCX │    │純文字  │    │結構MD  │    │6×JSON  │    │HTML 視 │
  │(機關分)│    │(類別) │    │SSOT⭐  │    │自動衍生│    │覺化介面│
  └────────┘    └────────┘    └────────┘    └────────┘    └────────┘
       ▲             ▲             ▲             ▲             │
       │ ◆_skip      │ 02_parse    │ 校對 +      │ 推斷邊      │
       │ ◆_manifest  │             │ LLM 重整    │ rate_lookup │
       │             │             │ 法源審查    │             │
       │       ┌─────┴─────────────┴─────┐               ┌─────▼─────┐
       │       │     輔助工具(_*.py)    │               │ GitHub    │
       └───────│  PDF 殘渣清整 / Title   │               │ Actions   │
               │  Summary / Tags / PII   │◀─────────────│ CI/CD     │
               │  Status / source_url    │               │ Pages     │
               └─────────────────────────┘               └───────────┘

  視圖層級:Landing → 情境檢索 → 條文庫(抽屜)→ 試算表 → 比較模式
  控管層級:DATA_VERSION 同步 / WCAG AA / 馬卡龍 token / 信度三層免責""")

    # 二、八階段詳細解構
    add_heading(d, "二、八階段流水線詳細解構", 1)

    stages = [
        {
            "no": "1",
            "name": "資料蒐集(00_source)",
            "input": "政府公開 PDF / DOCX / MD",
            "output": "依機關分類的原檔",
            "tool": "人工放檔 + _manifest.csv 補 metadata",
            "decisions": [
                "機關分流(01 國科會 / 02 成大 / 03 教育部 / 04 主計總處)",
                "檔名規則:保留原始名,過長者改 {法規名}_{識別}.ext",
                "_skip.txt 過濾雜訊(目次殘片、已存在 .md 對應 PDF)",
                "_manifest.csv:推斷不到時的 fallback metadata",
            ],
            "controls": [
                "00_source/ 唯讀,腳本不可寫入",
                ".gitignore 排除 00_source/(避免版權與肥 repo)",
                "敏感資訊禁止放入(repo 完全公開)",
            ],
            "risk": "原檔遺失 → 重建斷鏈",
        },
        {
            "no": "2",
            "name": "抽取(01_extract.py)",
            "input": "00_source 全檔",
            "output": "01_extracted/{類別}/*.txt + .meta.json",
            "tool": "pdfplumber(文字)/ paddleocr(掃描)/ python-docx",
            "decisions": [
                "文字型 vs 掃描型:預設文字型,需 OCR 才加 --ocr",
                "PARENT_KEYWORDS 母題推斷:特定法規名 > 母題本名 > 泛費目詞",
                "lazy import:重依賴只在需要時載入",
                "斷掉的 form-feed(\\f)在 strip_md_header 移除",
            ],
            "controls": [
                "_skip.txt 自動跳過",
                "費目名(如「支出憑證」)不單獨當關鍵字(避免誤判)",
                "退出碼:0 成功 / 1 環境錯 / 2 有錯誤",
            ],
            "risk": "母題推斷錯誤 → 後續分類全錯",
        },
        {
            "no": "3",
            "name": "解析(02_parse.py)",
            "input": "01_extracted/*.txt",
            "output": "02_markdown/*.md(草稿)",
            "tool": "Python 規則式 regex + 中文數字 ↔ 阿拉伯轉換",
            "decisions": [
                "front-matter schema(id/type/parent/title/tags/related/source/version)",
                "條文必須逐字保留(不正規化全形/不轉編號為 list)",
                "should_skip 安全網:已含 reviewed 不覆寫",
                "--force-reviewed 才能蓋已校對檔",
            ],
            "controls": [
                "schema 必填欄位由 04_validate.py 把關",
                "ID 編碼規則:{類別}-{母題}-{三位序號},刪除不重用",
                "MD 內禁止 HTML 與 <script>",
            ],
            "risk": "schema 變更需破壞性遷移所有 MD",
        },
        {
            "no": "4",
            "name": "校對(autoreview + 人工 + LLM 三層)",
            "input": "02_markdown 草稿",
            "output": "02_markdown 校對版(含 review_level)",
            "tool": "_batch_autoreview / _fix_titles v1+v2+_polish / _retitle_apply / _resummary_apply / 人工 / Claude Code subagent 16 batch",
            "decisions": [
                "review_level 三層:人工 / 自動初校 / llm精校",
                "Title 三段把關:截斷+檔名+TODO → 前言式語意模板 → 20 字硬上限+PII 重抽",
                "Summary 情境句型化:適用場景+核心規定+關鍵限制",
                "PII 清整:函稿表頭整行刪除 + 內嵌 email/電話遮蔽",
                "Tags 三步:同義詞合併 → 內容驅動補強 → 移除冗餘泛 tag",
                "Status 自動偵測「(刪除)」標 已廢止",
                "source_url 由子字串 mapping 表批次補(99.8% 涵蓋)",
            ],
            "controls": [
                "PHONE_RE 必須有分隔符(避免誤抓公文字號 0940006759)",
                "PII 子字串需涵蓋全形 + 半形冒號(法規 PDF 多用全形)",
                "_polish_titles 為最後一道:含 PII 標記或 >20 字必重抽",
                "已廢止 + effective_period 例外保留前端顯示",
                "summary_pending 由前端區分「真摘要 / 待補」",
            ],
            "risk": "LLM 重整 481 份「自動初校」實質仍是 LLM 而非人工",
        },
        {
            "no": "5",
            "name": "內容深化(scenarios + 法源審查)",
            "input": "校對完 MD + 情境設計",
            "output": "scenarios_manual.json + baseline_attachments.json",
            "tool": "_augment_scenarios / _audit_scenarios / _fix_stale_scenarios / 人工",
            "decisions": [
                "中立角色原則:不給判斷,每結論須有 A/B/C/D 法源",
                "scenarios 為展示用,非 SSOT,不進 03_index/",
                "情境卡 schema 12 欄(原 5 + 2026-05-01 加 7)",
                "6 個情境樹 root + sub_scenarios 整併(B1 一次到位)",
                "auto 卡 A1 決策停載(2026-05-01)",
                "deprecated 旗標:被 root flow 完整覆蓋的子卡",
            ],
            "controls": [
                "新增情境前必搜 A/B/C/D 全四分類(2026-05-02 教訓)",
                "每結論詞 / 數字 / 條件必對應原文,不可超出文義",
                "subtitle / caveats 用「條文規定...」中性化語氣",
                "違反原則的既有內容列待處理清單,不自行修",
                "可寫 _audit_scenario_sources.py 自動跑檢查",
            ],
            "risk": "AI 助手主動「常識補強」加結論 → 違反中立",
        },
        {
            "no": "6",
            "name": "索引建立(03_build_index.py)",
            "input": "02_markdown 全檔",
            "output": "03_index/(nodes / edges / tags / search_index / _meta / rate_lookup)",
            "tool": "Python 規則式 + body_plain 引用偵測",
            "decisions": [
                "推斷邊規則:「第 N 條/點」「QN」自動產出 _inferred",
                "跨母題推斷開放:前 40 字內出現另一母題名即跨",
                "rate_table 序列化:flat / sectioned / multiline / colspan",
                "rate_lookup 從 searchable: true 抽出",
                "extract_summary:前端 summary 取 100 字 + strip 自動初校標記",
            ],
            "controls": [
                "推斷邊已寫回 SSOT(425 條)→ 變人工邊(目前 482 全混雜)",
                "_meta.json 寫 last_indexed / node_count / status_counts",
                "certainty 序列化:explicit / inferred / contested",
            ],
            "risk": "推斷邊回寫 SSOT 後,再跑時新發現的才以 _inferred 出現",
        },
        {
            "no": "7",
            "name": "驗證(04_validate.py)",
            "input": "03_index/* + 02_markdown/*",
            "output": "驗證報告(stdout)",
            "tool": "Python 一致性檢查",
            "decisions": [
                "預設 / --strict 兩模式",
                "退出碼:0 / 1 / 2 / 3",
                "DoD:預設模式 exit 0,可有 warnings",
            ],
            "controls": [
                "schema 必填欄位、related 雙向一致、孤立節點警告",
                "CI 自動跑(GitHub Actions push to main)",
                "DATA_VERSION 與前端 query string 同步",
            ],
            "risk": "驗證項目仍偏 schema,語意正確性未檢",
        },
        {
            "no": "8",
            "name": "發布(04_web + GitHub Pages + CF)",
            "input": "03_index + 04_web/data/*",
            "output": "https://ntnick-web.github.io/gov-expense-kb/",
            "tool": "純 HTML/CSS/JS + GitHub Actions + Cloudflare Web Analytics",
            "decisions": [
                "純靜態無框架(禁 React/Vue/jQuery/CDN)",
                "馬卡龍 7 色 token + WCAG AA",
                "Landing 三入口卡(2026-04-30 加)",
                "lazy-render(條文 30 / 情境 3 sections)",
                "DATA_VERSION ?v=YYYY-MM-DDx 應用層快取破壞",
                "雙授權:程式 MIT / 內容 CC BY 4.0",
            ],
            "controls": [
                "CF Analytics token 嵌 index.html(無 cookie/PII)",
                ".nojekyll 必要(Pages 接收原檔)",
                "URL hash 進站不再觸發 view 切換(2026-04-27m)",
                "data-init-view 須在 switchView 開頭 remove(關鍵 fix)",
                "回滾路徑:index-old.html / git revert / mv 三選一",
            ],
            "risk": "single-file index.html 5424 行,維護門檻高",
        },
    ]

    for s in stages:
        add_heading(d, f"階段 {s['no']}:{s['name']}", 2)
        # 4 列 metadata 表
        add_table(
            d,
            ["項目", "內容"],
            [
                ["輸入", s["input"]],
                ["輸出", s["output"]],
                ["工具", s["tool"]],
                ["主要風險", s["risk"]],
            ],
            col_widths_cm=[2.5, 14.5],
            zebra=True,
        )
        add_para(d, "▍ 決策點", bold=True, color="3B6E8F", size=10.5)
        for dec in s["decisions"]:
            add_bullet(d, dec)
        add_para(d, "▍ 控管條件", bold=True, color="C84F4F", size=10.5)
        for ctrl in s["controls"]:
            add_bullet(d, ctrl, color="6A4040")
        d.add_paragraph()

    # 三、決策點 / 控管條件矩陣
    d.add_page_break()
    add_heading(d, "三、決策點 × 控管條件 矩陣表", 1)
    add_para(d, "全 8 階段彙整為「核心決策點」「控管邊界」「失效模式」三維對照,作為新進維護者的速查清單。")

    matrix_rows = [
        ["全域", "中立角色原則", "每結論需有真實法源(A/B/C/D)", "AI 主動「常識補強」(已記錄 5 件)", "P0"],
        ["全域", "法源位階", "A 核心 > B 標準 > C 解釋 > D 問答", "僅以 D 推翻 A 文義", "P0"],
        ["全域", "SSOT 原則", "02_markdown 唯一事實源", "改前端不改 MD → 索引漂移", "P0"],
        ["階段 1", "_skip.txt", "排除已存在 .md 的重複 PDF", "未維護 → 重複抽取雜訊", "P1"],
        ["階段 2", "PARENT_KEYWORDS", "特定法規名優先,費目名不獨用", "費目名單獨當 keyword → 誤判", "P1"],
        ["階段 3", "reviewed 安全網", "已校對檔不被覆寫", "誤用 --force-reviewed 蓋寫", "P1"],
        ["階段 4", "title ≤ 20 字", "_polish_titles 最後一道", "規則衝突或 PII 滲入", "P1"],
        ["階段 4", "PHONE_RE 分隔符", "避免抓公文字號", "改寫 regex 漏掉條件", "P2"],
        ["階段 4", "PII 全形冒號", ": 與:都涵蓋", "regex char class 缺一", "P1"],
        ["階段 5", "A/B/C/D 全分類搜", "新情境法源檢查 SOP", "只搜 A/B/D 漏 C(2026-05-02 教訓)", "P0"],
        ["階段 5", "primary_ids 真實性", "節點必存在 nodes.json", "ID typo / 已刪 ID", "P0"],
        ["階段 5", "中性化語氣", "用「條文規定」取代「應該」", "AI 寫成判斷句", "P1"],
        ["階段 6", "推斷邊跨母題", "前 40 字內含母題名才跨", "規則放寬 → 假連結爆量", "P2"],
        ["階段 6", "rate_table searchable", "opt-in,需顯式宣告", "未宣告 → Ctrl+K 找不到", "P2"],
        ["階段 7", "DoD exit 0", "預設模式無 errors", "CI 失敗未及時修", "P1"],
        ["階段 7", "DATA_VERSION 同步", "MD 改 + 前端 ?v= 同升", "舊 JSON 持續快取", "P1"],
        ["階段 8", "data-init-view fix", "switchView 開頭 remove", "view 切換 mobile leak", "P1"],
        ["階段 8", "回滾三路徑", "git revert / index-old / mv", "新版上線即破未保留路徑", "P0"],
    ]
    add_table(
        d,
        ["階段", "決策點 / 控管", "判準", "失效模式", "優先級"],
        matrix_rows,
        col_widths_cm=[1.8, 3.6, 4.8, 4.8, 1.5],
    )

    # 四、核心邏輯與護城河
    d.add_page_break()
    add_heading(d, "四、核心邏輯與護城河分析", 1)

    add_heading(d, "4.1 核心思想三主軸", 2)
    add_para(d, "1. 中立角色 — 不下判斷、不主動補實務知識,只引用條文。所有結論都可回溯到法源原文,使用者自行決定核銷可否。", bold=True)
    add_para(d, "2. 單一事實來源 — 02_markdown 是 SSOT,03_index 與 04_web 皆可從 SSOT 完整重建。即使前端崩潰,MD 仍是完整可讀的法規知識庫。", bold=True)
    add_para(d, "3. 零依賴可離線 — 純靜態無框架,GitHub Pages 託管即可運作;Python 標準庫 + 三件 PDF 處理套件即可從零重建。", bold=True)

    add_heading(d, "4.2 護城河(競品難以快速複製)", 2)
    moat_rows = [
        ["法源完整性", "520 節點覆蓋 A/B/C/D 四類,涵蓋率近 100%", "需法務協作 + 主計總處公開資料整理 6 個月以上"],
        ["中立角色 SOP", "明確法源位階 + 5 件處理紀錄 + 17 章審查清單", "等同把法規顧問的審慎性內建為流程"],
        ["馬卡龍設計系統", "7 色 token + WCAG AA + 深色模式對應", "市面政府網站普遍硬色彩,品牌辨識度高"],
        ["結構化費率表", "8 張 B 類標準表 + Ctrl+K 直接答案 + 城市 fallback", "需 schema 設計 + 全文搜尋整合,獨家功能"],
        ["保險試算 widget", "lookup_type:insurance 條件式 widget", "lookup_type 模式可外推到其他費率類型"],
        ["LLM 16 batch 模式", "免費 subagent 大規模 retitle/resummary 範式", "需深度懂 Claude Code subagent 機制"],
        ["雙模式關聯邊", "人工 cites/explains/answers + 推斷 _inferred", "需 body_plain 引用偵測 + 跨母題規則"],
        ["信度三層免責", "explicit/inferred/contested + 三層 UI 提示", "比一般「免責頁」更精細,逐節點標記"],
        ["靜態網站 + CF Analytics", "零後端 / 無 cookie / GDPR 友善", "符合政府網站對 PII 的低風險偏好"],
    ]
    add_table(
        d,
        ["護城河面向", "本資料庫做法", "為何難複製"],
        moat_rows,
        col_widths_cm=[3.0, 6.5, 6.5],
        header_bg="2E5B7A",
    )

    add_heading(d, "4.3 SWOT 分析", 2)

    add_para(d, "S 優勢", bold=True, color="2E7D4F", size=11.5)
    for x in [
        "單向資料流 + 可重建衍生物,任一階段可獨立檢視與除錯",
        "MD + JSON 雙層,Git diff 友善、版控歷史清楚",
        "純靜態,部署成本近零,任何地方可一鍵復現",
        "中立原則明確,法律風險低",
        "已建立 5 件法源審查處理紀錄與 SOP,維護紀律可學習",
    ]:
        add_bullet(d, x, color="2E7D4F")

    add_para(d, "W 劣勢", bold=True, color="C84F4F", size=11.5)
    for x in [
        "13+ 個 _*.py 一次性腳本散落,無 _common.py 抽取,維護心智負擔大",
        "review_level 481 份(>92%)為「自動初校」,人工複核率僅約 7%",
        "推斷邊已寫回 SSOT,失去區分「人工 / 推斷」的能力",
        "scenarios_manual.json 單檔 2973 行,跨情境變更難 diff",
        "index.html 5424 行單檔,inline JS+CSS,維護門檻高",
        "母題僅 3 個,跨域擴充慢(酬勞費已規劃但未啟動)",
        "67/73 情境卡尚未補 caveats / example / template(內容深化進度 8%)",
    ]:
        add_bullet(d, x, color="6A4040")

    add_para(d, "O 機會", bold=True, color="3B6E8F", size=11.5)
    for x in [
        "酬勞費母題(講座 / 出席 / 稿費)— 預估 18-25 卡,4-6 週可加",
        "後續 5 母題:共通性費用 / 加班費 / 公務車輛 / 教育部 / 國科會",
        "rate_table lookup_type 模式可外推:旅費表、加班費分級表",
        "Decision tree 從 9 → ~30 配合酬勞費同步擴增",
        "CF Workers + D1 後端事件追蹤(scenario 點擊 / search query)",
        "整合 LLM 為「閱讀引導員」(不下判斷,僅引導使用者翻條文)",
    ]:
        add_bullet(d, x, color="3B6E8F")

    add_para(d, "T 威脅", bold=True, color="8B5A00", size=11.5)
    for x in [
        "主計總處法規修訂 → 大量條文需重抽 + LLM 重整(2024 ~ 2026 已換過 1 輪)",
        "AI 助手「主動補強」未被察覺 → 違反中立角色(已 1 次教訓)",
        "GitHub Pages CDN 邊緣快取 ~10 分鐘 → 緊急修補延遲",
        "主計總處 URL 結構變動 → source_url 連結失效",
        "節點數突破 1000 → vanilla substring 搜尋失效,須 FlexSearch 中文分詞",
        "酬勞費母題的「政務官 vs 一般機關」分流可能與既有 schema 衝突",
    ]:
        add_bullet(d, x, color="6A4A20")

    # 五、改善建議
    d.add_page_break()
    add_heading(d, "五、改善建議(短中長期 28 項)", 1)

    add_heading(d, "5.1 短期(1-2 週,可立即動手)", 2)
    short_term = [
        ["#1", "抽 _common.py 共用模組",
         "split_fm / render_fm / extract_section / SYNONYMS / PHONE_RE 等抽出。13 腳本各自局部複製,門檻已過。",
         "降低 50% 重複代碼;規則變更只改一處"],
        ["#2", "拆 04_web/index.html 成 ESM module",
         "5424 行單檔 → app/main.js + ui/{drawer,scenarios,library}.js + design/tokens.css。",
         "提升維護性;diff 可讀;支援 IDE 跳轉"],
        ["#3", "下架 v2 / v3 / old 三舊版",
         "index-v2.html(載 scenarios.json,落後)/ v3 / old + assets/{app.js,style.css} 做封存後刪。",
         "減少新進者誤改錯版本"],
        ["#4", "_audit_scenario_sources.py 自動跑法源檢查",
         "scenarios_manual.json 內每張卡的 primary_ids → 比對 nodes.json + 比對 summary 是否含結論關鍵詞",
         "預防 2026-05-02 那種「漏搜 C 類」教訓重演"],
        ["#5", "review_level 拆出「LLM 精校」",
         "目前 481 份「自動初校」實質為 LLM。新增 llm精校 標記後重跑 _migrate_review_level.py",
         "讓使用者能看出真正人工 vs LLM 比例"],
        ["#6", "DATA_VERSION 自動同步",
         "GitHub Actions 在 push to main 時自動把 ?v=YYYY-MM-DDx 寫入 index.html(讀 _meta.last_indexed)",
         "杜絕忘記升版造成的舊快取"],
        ["#7", "刪 nodes_sample.json / _retitle_proposals/ / _resummary_proposals/",
         "untracked dev artefact ~2 MB,已 apply。詳見『資料夾整理報告』",
         "釋出磁碟 + 減少 git status 雜訊"],
        ["#8", "更新 README.md 與 docs/01_architecture.md",
         "目前仍寫 513 節點 + 3 視圖,實際 520 節點 + 6 視圖(landing/情境/條文庫/試算/抽屜/比較)",
         "對外資訊一致;新進者不會被誤導"],
        ["#9", "CLAUDE.md 拆檔",
         "92 KB / 700 行單檔 → 核心原則(< 30K)+ docs/changelog.md(歷史變更)+ docs/_review_log.md(法源審查)",
         "降低 AI 助手每次閱讀成本;歷史紀錄保留但不阻塞"],
        ["#10", "scenarios JSON Schema validate",
         "JSON Schema 定義 12 欄位 → CI 自動驗 caveats[].severity ∈ {stop,warn,info},template[].type 等",
         "schema 漂移即時 fail"],
    ]
    add_table(d, ["#", "建議", "做法", "預期效益"], short_term,
              col_widths_cm=[0.8, 3.5, 8.0, 4.7])

    add_heading(d, "5.2 中期(1-2 個月,需設計)", 2)
    mid_term = [
        ["#11", "推斷邊獨立追蹤",
         "新欄位 related_inferred 與 related 並列;_write_inferred_related 改寫 related_inferred",
         "可重建「人工 vs 推斷」之分,信度提升"],
        ["#12", "scenarios 拆按母題切多檔",
         "scenarios/{domestic,abroad,voucher,honorarium}.json + build 時合併,類似 baseline_attachments 模式",
         "diff 可讀;變更影響面清楚"],
        ["#13", "67 卡內容深化(caveats / example / template)",
         "用 16 batch subagent 範式逐母題重整 — 套用 retitle/resummary 同樣手法",
         "情境卡完整度從 8% → 80%"],
        ["#14", "Decision tree 從 9 → 30",
         "配合酬勞費母題同步補,每 root 內增 3-5 個 inner flow",
         "進階使用者引導路徑完整"],
        ["#15", "rate_lookup 加「比照鄰國」推斷",
         "未列載國家 → 顯示「比照 X 國 → 其他 → Y 美元」+ 標警示",
         "解決政府附註 §2 已知缺口"],
        ["#16", "增量 build_index",
         "從 mtime 偵測變動的 MD,只重 index 受影響子集,而非全量",
         "5x build 速度提升,鼓勵頻繁迭代"],
        ["#17", "CI 加 link-check + JSON Schema",
         "驗 source_url 200 / scenarios 結構 / 內部 anchor 對得上",
         "上游 URL 失效自動報警"],
        ["#18", "保險表 widget 抽 lookup_type 通用化",
         "lookup_type: insurance / range_table / monthly_to_daily 三類,共用 wireWidget runtime",
         "未來新費率表 widget 不必重寫"],
        ["#19", "同義詞 21 → 50 組",
         "從前端搜尋 0 命中 query log 補 — 需先加 query log",
         "搜尋成功率 +10%(估)"],
        ["#20", "test_smoke.py + Playwright headless",
         "GitHub Actions 跑 5 條金路徑(載入 / 切 view / 開抽屜 / 試算 / Ctrl+K)",
         "前端 regression 自動發現"],
    ]
    add_table(d, ["#", "建議", "做法", "預期效益"], mid_term,
              col_widths_cm=[0.8, 3.5, 8.0, 4.7])

    add_heading(d, "5.3 長期(3-6 個月,戰略級)", 2)
    long_term = [
        ["#21", "酬勞費母題上線",
         "P3 規劃,18-25 卡,法源:中央政府各機關學校出席費及稿費支給要點 + 講座鐘點費支給表",
         "覆蓋第 4 大常用核銷情境;議價能力 +1"],
        ["#22", "後續 5 母題擴充",
         "順序:#2 共通性費用 → #3 加班費 → #4 公務車輛 → #5 教育部 → #6 國科會",
         "從旅費庫升級為「政府支出全領域」"],
        ["#23", "條文版本歷史(law_version_history)",
         "每條文記錄修法日 + 變更摘要 + 替代節點 ID,前端時間軸顯示",
         "對應主計總處每年修法週期;審計查歷史可追溯"],
        ["#24", "FlexSearch 中文分詞",
         "節點 > 1000 觸發。配合同義詞表訓練分詞器,前端搜尋速度 +5x",
         "中文搜尋體驗提升;多詞 query 精度 +20%"],
        ["#25", "後端事件追蹤(CF Workers + D1)",
         "scenario 點擊 / drawer 開啟 / search query → DAU/MAU 後決定付費功能",
         "為商業化打基礎"],
        ["#26", "LLM「閱讀引導員」整合",
         "嚴守中立 — 不下判斷,只引導翻條文。Side-panel chat 式,引用必附條文 ID",
         "降低首次使用者學習成本;不違中立原則"],
        ["#27", "離線桌面版(Tauri)",
         "Rust + WebView 包成 .exe / .dmg / .deb,完全離線可用",
         "符合政府機關內網場景;敏感案例可離線查"],
        ["#28", "API 開放(read-only)",
         "GitHub Pages 上既有 03_index/*.json 已是 API。寫 docs/api.md 對外開放",
         "讓第三方工具(LINE bot、Slack、本機 CLI)接入"],
    ]
    add_table(d, ["#", "建議", "做法", "預期效益"], long_term,
              col_widths_cm=[0.8, 3.5, 8.0, 4.7])

    # 六、優化決策樹
    d.add_page_break()
    add_heading(d, "六、優化排序決策樹", 1)
    add_para(d, "在 28 項建議中,該先做哪一個?以下為「投入產出比」排序決策樹(由上而下優先)。")
    add_ascii_block(d, """\
                            ┌──────────────────┐
                            │ 接下來 1 週要做? │
                            └────────┬─────────┘
                                     │
            ┌────────────────────────┼────────────────────────┐
            ▼                        ▼                        ▼
  【清地基】立即執行         【補護欄】1-2 週         【拓疆土】1-3 月
   #7 刪 dev artefact         #4 法源檢查腳本          #21 酬勞費母題
   #3 下架 v2/v3/old          #1 抽 _common.py         #13 67 卡內容深化
   #8 更新 README             #5 review_level 拆       #11 推斷邊獨立
   #9 CLAUDE.md 拆檔          #6 DATA_VERSION 自動同步 #12 scenarios 拆檔
   (~2-3 天 / 高效益)        (~1 週 / 中效益)         (~1-3 月 / 戰略效益)
            │                        │                        │
            └────────────────────────┴────────────────────────┘
                                     │
                              全做完?↓ 否則暫停
                                     │
                          ┌──────────┴──────────┐
                          ▼                     ▼
                   【系統化】3-6 月       【商業化】6-12 月
                    #2 拆 ESM module       #25 後端事件追蹤
                    #16 增量 build         #26 LLM 引導員
                    #20 Playwright e2e     #27 離線桌面版
                    #24 FlexSearch         #28 API 開放""")

    add_para(d, "建議路徑", bold=True, color="2E5B7A", size=11.5)
    add_numbered(d, "本週啟動「清地基」#7 #3 #8 #9 — 純收益,無風險。", 1)
    add_numbered(d, "下週啟動「補護欄」#4 #1 #5 #6 — 寫一次,長期受益。", 2)
    add_numbered(d, "5 月啟動「拓疆土」#21 酬勞費母題 — 此為下次內容深化主推項。", 3)
    add_numbered(d, "8 月後再評估「系統化」與「商業化」 — DAU/MAU 數據夠才動。", 4)

    # 結語
    add_heading(d, "結語", 1)
    add_para(d, "本資料庫架構在「政府支出法規」這個冷門但剛性需求的領域,已建立一套完整、可重建、易維護的工程紀律。最大的護城河不是技術,而是「中立角色 + 法源位階 + 校對 SOP」三位一體的內容紀律 — 這點是 AI/搜尋引擎短期內無法複製的。")
    add_para(d, "短期應務實清地基,把已過期的多版本檔、未抽出的共用模組、尚未拆檔的單體 index.html 處理完;中期在內容深化與工具化下功夫;長期再以酬勞費母題開啟商業化路徑。")
    add_para(d, "─── 報告完 ───", align=WD_ALIGN_PARAGRAPH.CENTER, color="888888", italic=True)

    d.save(out_path)
    print(f"  ✓ {out_path.relative_to(ROOT)}")


# ─────────────────────────────────────────────────────────────
# 報告 2:資料庫資料夾整理報告
# ─────────────────────────────────────────────────────────────

def build_cleanup_report(out_path: Path):
    d = doc_init()

    # 封面
    add_heading(d, "資料庫資料夾整理報告", 0)
    add_para(d, "── 政府支出法規知識庫 ──", size=12, color="6A7B8C", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(d, "識別冗餘 / 一次性 / 過期檔案 + 提供分批清理 SOP",
             size=10.5, italic=True, color="6A7B8C", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(d, "報告日期:2026-05-01    掃描範圍:整個專案根目錄    建議釋出:約 2 MB(本機) + ~600 KB(repo)",
             size=9.5, color="888888", align=WD_ALIGN_PARAGRAPH.CENTER)
    d.add_paragraph()

    # 執行摘要
    add_heading(d, "執行摘要", 1)
    add_para(d, "整個專案根目錄掃描後,可分為四類處置:")
    add_bullet(d, "✅ 核心檔(必留):520 節點 MD、index 索引、主 index.html、要 docs、.github、LICENSE/README", color="2E7D4F")
    add_bullet(d, "🟧 dev artefact(本機可刪):2 MB 左右一次性 LLM 重整輸出 + sample JSON + audit txt", color="C8761F")
    add_bullet(d, "🟥 多版本舊檔(應從 git 退役):index-v2/v3/old + assets/* + scenarios.json + scenarios_auto.json + design-preview/ ≈ 600 KB", color="C84F4F")
    add_bullet(d, "🔵 文件待重構:CLAUDE.md(92 KB / 太多歷史)+ docs/_handoff_*.md(已用畢)+ README/01_architecture(過時)", color="3B6E8F")
    add_para(d, "建議分三批執行,每批可獨立回滾。詳見第三章。", bold=True)

    # 一、現況統計
    add_heading(d, "一、現況統計", 1)
    add_table(d, ["資料夾 / 檔", "大小", "Git 狀態", "用途", "處置"], [
        ["00_source/", "4.8 MB", "gitignored", "原始 PDF/DOCX(682 檔)", "✅ 本機保留"],
        ["01_extracted/", "5.1 MB(空目錄)", "gitignored", "抽取產物(0 檔)", "✅ 可全清,需要時重跑"],
        ["02_markdown/", "2.6 MB", "tracked", "SSOT(520 MD)", "✅ 核心,絕對保留"],
        ["03_index/", "1.6 MB", "tracked", "JSON 索引(6 檔)", "✅ 核心,自動產生"],
        ["04_web/", "1.2 MB", "tracked", "前端(含舊版)", "🟥 部分舊版可退"],
        ["05_scripts/", "2.7 MB", "部分 tracked", "腳本 + 一次性產物", "🟧 dev artefact 可清"],
        ["docs/", "132 KB", "tracked", "規格 + 法務 + 報告", "🔵 部分需重構"],
        ["design_bundle_WHBnmmq5/", "172 KB", "gitignored", "設計交接包", "✅ 本機保留(已 ignored)"],
        [".git/", "23 MB", "—", "版控歷史", "✅ 不動"],
        ["CLAUDE.md", "92 KB", "gitignored", "AI 助手常駐指引", "🔵 拆檔重構"],
        ["README.md", "6.4 KB", "tracked", "對外說明", "🔵 內容過時待更"],
        ["LICENSE.md", "3.3 KB", "tracked", "雙授權", "✅ 保留"],
        ["index.html(根)", "528 B", "tracked", "meta-refresh 重定向", "✅ 保留"],
        ["requirements.txt", "527 B", "tracked", "Python 依賴", "✅ 保留"],
        [".gitignore", "732 B", "tracked", "已寫好排除規則", "✅ 保留"],
        [".nojekyll", "0 B", "tracked", "Pages 必要", "✅ 保留"],
    ], col_widths_cm=[4.5, 2.5, 2.0, 4.5, 3.5])

    # 二、清理三批
    d.add_page_break()
    add_heading(d, "二、清理對象詳細清單(三批)", 1)

    # 第一批
    add_heading(d, "第一批:本機可立即刪(untracked dev artefact)", 2)
    add_callout(d,
                "🟧 釋出空間:約 2.0 MB    Git 影響:0    風險:極低(全為一次性產物)",
                "這批檔案皆為 LLM 大規模重整的產物,已 apply 到 02_markdown/。以 _開頭 已被慣例視為 dev,且在 git status untracked 區。可直接刪除,要回滾去 git log 查 commit 即可重做。",
                bg="FFF7E6", border="E8B85E")

    batch1 = [
        ["05_scripts/_retitle_proposals/", "924 KB", "untracked",
         "2026-04-29 LLM retitle 16 batch subagent 的輸入(inputs/)/ 輸出(outputs/)/ manifest.json / quality_report.txt / review.csv 全套留檔",
         "已透過 _retitle_apply.py 套到 02_markdown(460 張改),產物已不再使用"],
        ["05_scripts/_resummary_proposals/", "1.1 MB", "untracked",
         "同上,2026-04-29 LLM resummary 16 batch 的全套輸入輸出",
         "已透過 _resummary_apply.py 套到 02_markdown(479 張改),產物已不再使用"],
        ["05_scripts/__pycache__/", "32 KB", ".gitignore 已含",
         "Python bytecode 快取",
         ".gitignore 已寫 __pycache__/,可直接 rm -rf"],
        ["05_scripts/_certainty_review.csv", "—", "untracked",
         "_mark_certainty.py 的 review CSV(519 節點分布 explicit/inferred/contested)",
         "已 apply 到 MD,留底意義小"],
        ["05_scripts/_scenario_audit.txt", "—", "untracked",
         "_audit_scenarios.py 輸出(2026-04-29 stale scenarios 清單)",
         "已透過 _fix_stale_scenarios.py 修補,留底意義小"],
        ["04_web/data/nodes_sample.json", "9822 行", "untracked",
         "疑似前端開發過程的取樣 dump",
         "未被 index.html 引用(grep 0 命中);明顯為 debug 暫存"],
    ]
    add_table(d, ["路徑", "大小", "Git", "原本用途", "為何可刪"], batch1,
              col_widths_cm=[5.0, 1.5, 1.8, 5.5, 3.2])

    add_para(d, "建議執行(本機)", bold=True, color="2E7D4F")
    add_ascii_block(d, """\
# bash 一鍵清(可直接複製到專案根目錄執行)
rm -rf 05_scripts/_retitle_proposals/
rm -rf 05_scripts/_resummary_proposals/
rm -rf 05_scripts/__pycache__/
rm    05_scripts/_certainty_review.csv
rm    05_scripts/_scenario_audit.txt
rm    04_web/data/nodes_sample.json

# 結果驗證
git status                 # 應該全乾淨,untracked 區清空
du -sh 05_scripts/         # 從 2.7M 降到約 600K""")

    # 第二批
    d.add_page_break()
    add_heading(d, "第二批:應從 git 退役 + gitignore", 2)
    add_callout(d,
                "🟥 釋出空間:約 600 KB(repo)+ 維護心智成本大幅下降    Git 影響:有(需 commit)    風險:中(需確認回滾路徑)",
                "這批是「過渡期保留多版本」的舊檔,目前同時 tracked 於 git,但實際上已被新版正式版取代。CLAUDE.md §15 雖說『保留供 bookmark 不破』,但 (1) v2 還在載 scenarios.json(過時) (2) v3 是 prototype 早併入 (3) old 是凍結備份。建議改用 git tag 標記版本即可,而非保留實檔。",
                bg="FFE6E6", border="E08585")

    batch2 = [
        ["04_web/index-v2.html", "2430 行 / 128 KB", "tracked",
         "舊正式版(2026-04-28r 之前的 v2 設計版)",
         "✗ 仍 fetch scenarios.json(舊 monolith)",
         "git tag v2-archive + 刪檔"],
        ["04_web/index-v3.html", "1841 行 / 96 KB", "tracked",
         "V3 prototype(commit 6dfea70)",
         "✗ 馬卡龍重構已併入 index.html(0beb8af)",
         "git tag v3-prototype + 刪檔"],
        ["04_web/index-old.html", "235 行 / 12 KB", "tracked",
         "舊版凍結備份(母題泡泡圖架構)",
         "✗ 真要 fallback 用 git revert 比保留實檔乾淨",
         "git tag old-bubblechart + 刪檔"],
        ["04_web/assets/app.js", "150 KB", "tracked",
         "舊版 JS",
         "✗ 只有 index-old.html 引用",
         "隨 index-old 一起退"],
        ["04_web/assets/style.css", "—", "tracked",
         "舊版 CSS",
         "✗ 同上",
         "隨 index-old 一起退"],
        ["04_web/data/scenarios.json", "2684 行 / 100 KB", "tracked",
         "舊 monolith(74 manual + 43 auto 合併)",
         "✗ A1 決策已停載 auto;index.html 載 scenarios_manual",
         "改 fallback 來源 → 改 index.html → 刪 JSON"],
        ["04_web/data/scenarios_auto.json", "736 行 / 28 KB", "tracked",
         "auto 卡規則式生成 43 張",
         "✗ A1 決策已停載(2026-05-01)",
         "git tag auto-archive + 刪檔"],
        ["04_web/design-preview/", "5 檔", "tracked",
         "React + JSX 設計試做版",
         "△ index.html 還有連結「🎨 設計試做」",
         "若不想砍,移到獨立分支 design-experiment"],
    ]
    add_table(d, ["路徑", "大小", "Git", "原本用途", "為何退役", "建議做法"],
              batch2, col_widths_cm=[3.8, 2.0, 1.4, 2.8, 3.3, 3.7])

    add_para(d, "建議執行步驟", bold=True, color="C84F4F")
    add_numbered(d, "先打 git tag 留歷史錨點:git tag v2-archive 8aaf079; git tag v3-prototype 6dfea70; git tag old-bubblechart <commit>", 1)
    add_numbered(d, "改 04_web/index.html 第 2464 行,刪 scenarios.json fallback,讓 scenarios_manual.json 為唯一來源", 2)
    add_numbered(d, "git rm 04_web/index-v2.html 04_web/index-v3.html 04_web/index-old.html 04_web/assets/app.js 04_web/assets/style.css 04_web/data/scenarios.json 04_web/data/scenarios_auto.json", 3)
    add_numbered(d, "決議 design-preview/:① 移到 git branch design-experiment 後從 main 刪 ② 留著但加 noindex,選一", 4)
    add_numbered(d, "更新 CLAUDE.md §15 與 README:刪除「v2 / v3 / old 並存」段落,改寫『歷史版本以 git tag 取得』", 5)
    add_numbered(d, "提一個 PR:chore(cleanup): 退役多版本舊檔,以 git tag 取代", 6)

    # 第三批
    d.add_page_break()
    add_heading(d, "第三批:CLAUDE.md / docs/ 重構建議", 2)
    add_callout(d,
                "🔵 釋出空間:約 60 KB(CLAUDE.md 拆瘦)    Git 影響:有(需 commit + AI 助手重新適應)    風險:低(純文件搬家)",
                "CLAUDE.md 92 KB / 700 行已過於肥大,內含:核心原則、最新狀態、實作筆記、歷史回滾、法源審查紀錄、改善計畫等多種異質內容。建議拆成『常駐指引(短)』+『歷史紀錄(可分檔)』,讓 AI 助手每次讀的成本更低、更新成本也更低。",
                bg="E6F2FB", border="6BA8D9")

    add_heading(d, "3.1 CLAUDE.md 拆檔建議", 3)
    add_table(d, ["原章節", "目前位置", "建議去處", "理由"], [
        ["核心原則(中立角色 / 法源位階)", "§0 + §17", "保留 CLAUDE.md(濃縮版)",
         "高頻使用,AI 每次都需讀"],
        ["最新狀態速覽", "§0", "保留(壓縮到 < 30 行)",
         "高頻使用"],
        ["技術棧 / 資料夾規範 / ID 規則", "§2-§4", "保留(已濃縮)",
         "高頻使用"],
        ["MD schema / 視圖架構 / 編碼規範", "§5-§7", "保留",
         "中頻使用"],
        ["禁止事項 + DoD", "§8-§9", "保留",
         "高頻使用"],
        ["標準管線 + 工作流程", "§10-§11", "保留",
         "高頻使用"],
        ["重要決策紀錄(2026-04-25 採 MD)", "§12 引導", "已在 docs/decisions.md",
         "保留 §12 一句話引導"],
        ["求助時機 + 已知擴充點", "§13-§14", "保留",
         "中頻使用"],
        ["線上部署 §15 + v2 重構 + 馬卡龍 + 整併深化 v3", "§15", "→ docs/changelog.md(新建)",
         "歷史筆記,寫一次後不再讀;留 §15 部署 SOP 即可"],
        ["對外開放 / 授權 / 隱私", "§16", "保留(濃縮)",
         "中頻使用,涉及法務"],
        ["法源審查紀錄 §17", "§17", "→ docs/_review_log.md(新建)",
         "持續累積的審查紀錄,單獨檔好維護"],
    ], col_widths_cm=[4.8, 2.5, 4.5, 5.0])

    add_para(d, "預估拆完後 CLAUDE.md 約剩 25-30 KB(從 92 KB)。", bold=True, color="3B6E8F")

    add_heading(d, "3.2 docs/ 內容檢視", 3)
    add_table(d, ["檔案", "用途", "現況", "建議"], [
        ["01_architecture.md", "系統架構", "✗ 過時(寫 3 視圖、總覽圖未含 landing)", "更新至 6 視圖 + landing"],
        ["02_data_schema.md", "schema 規格", "△ scenarios 7 新欄位未補", "補 caveats / example / template / baseline / flow_root / sub_scenarios / deprecated"],
        ["03_id_convention.md", "ID 規則", "✓ 同步", "保留"],
        ["04_ui_spec.md", "UI 規格", "△ 需確認是否同步當前 v2 設計", "與當前 index.html 對齊"],
        ["05_workflow.md", "處理 SOP", "✓ 同步", "保留(可補 LLM 16 batch 段)"],
        ["06_tags_taxonomy.md", "標籤分類", "✓ 同步", "保留"],
        ["about.md / privacy.md / terms.md", "法務(footer 露出)", "✓ 同步", "保留"],
        ["decisions.md", "ADR", "✓ 持續累積", "保留(可補 2026-04-29 ~ 2026-05-01 三筆)"],
        ["_handoff_optimization.md", "2026-04-29 LLM 重整 handoff", "✗ 已用畢", "→ docs/_archive/(封存)"],
        ["_handoff_retitle.md", "同上,retitle 專屬", "✗ 已用畢", "→ docs/_archive/(封存)"],
        ["_未來優化規劃報告_2026-05-01.docx", "2026-05-01 規劃文件", "✓ 仍用", "保留"],
    ], col_widths_cm=[5.0, 4.0, 3.5, 4.3])

    add_heading(d, "3.3 README.md 檢視", 3)
    add_para(d, "目前 README.md 第 5 行仍寫:「目前資料規模:513 節點(...)」「三視圖(泡泡概覽 / 條文庫 / 關聯圖)」。",
             color="6A4040")
    add_para(d, "實際:520 節點、6 視圖(landing / 情境檢索 / 條文庫 / 試算表 / 抽屜 / 比較模式)。",
             bold=True, color="2E7D4F")
    add_para(d, "建議:更新節點數(改用 _meta.json 動態插入語法)+ 視圖描述。", color="3B6E8F")

    # 三、執行 SOP
    d.add_page_break()
    add_heading(d, "三、執行 SOP(三週分批)", 1)

    add_heading(d, "週 1:第一批清理(本機)", 2)
    add_para(d, "目標:釋出 ~2 MB 本機空間 + 清空 git status untracked 區", bold=True)
    add_ascii_block(d, """\
□ Day 1  確認所有產物已 apply(對 02_markdown/ 抽 5 個樣本檢查 review_level 與 title)
□ Day 1  rm -rf 05_scripts/_retitle_proposals/
□ Day 1  rm -rf 05_scripts/_resummary_proposals/
□ Day 1  rm -rf 05_scripts/__pycache__/
□ Day 1  rm 05_scripts/_certainty_review.csv 05_scripts/_scenario_audit.txt
□ Day 1  rm 04_web/data/nodes_sample.json
□ Day 1  git status 確認全乾淨
□ Day 2  在 .gitignore 補 _retitle_proposals/、_resummary_proposals/、_*.csv、_*.txt(以防萬一)""")

    add_heading(d, "週 2:第二批退役(repo)", 2)
    add_para(d, "目標:減少 ~600 KB tracked 體積 + 心智負擔下降", bold=True)
    add_ascii_block(d, """\
□ Day 1  git tag v2-archive 8aaf079
□ Day 1  git tag v3-prototype 6dfea70
□ Day 1  git tag old-bubblechart $(git log --oneline -- 04_web/index-old.html | tail -1 | awk '{print $1}')
□ Day 2  改 04_web/index.html 第 2464 行,刪 scenarios.json fallback
□ Day 2  Ctrl+F 全 repo 確認沒有其他地方還引用 scenarios.json / scenarios_auto.json
□ Day 3  git rm 04_web/index-v2.html
□ Day 3  git rm 04_web/index-v3.html
□ Day 3  git rm 04_web/index-old.html
□ Day 3  git rm 04_web/assets/app.js 04_web/assets/style.css
□ Day 3  git rm 04_web/data/scenarios.json 04_web/data/scenarios_auto.json
□ Day 4  決議 design-preview/(① 砍 ② 移分支 ③ 留 + noindex)
□ Day 5  本機跑 python -m http.server,測試 5 條金路徑(landing / 情境 / 條文 / 試算 / Ctrl+K)
□ Day 5  測試 https://ntnick-web.github.io/gov-expense-kb/?v=YYYY-MM-DDx 強制重整
□ Day 6  PR: chore(cleanup): 退役多版本舊檔,以 git tag 取代""")

    add_heading(d, "週 3:第三批文件重構", 2)
    add_para(d, "目標:CLAUDE.md 從 92 KB → 25-30 KB,docs 重新對齊現況", bold=True)
    add_ascii_block(d, """\
□ Day 1  建立 docs/changelog.md(從 CLAUDE.md §15 整段搬過去 + 加日期錨點)
□ Day 1  建立 docs/_review_log.md(從 CLAUDE.md §17 整段搬過去)
□ Day 2  CLAUDE.md 拆檔:刪 §15 v2 重構 / 馬卡龍 / 整併深化 v3 三大段
□ Day 2  CLAUDE.md §0 最新狀態速覽壓縮到 30 行內
□ Day 3  更新 README.md(節點數 / 視圖描述)
□ Day 3  更新 docs/01_architecture.md(總覽圖加 landing + 6 視圖)
□ Day 4  更新 docs/02_data_schema.md(scenarios 7 新欄位)
□ Day 4  建立 docs/_archive/ 移入 _handoff_optimization.md / _handoff_retitle.md
□ Day 5  更新 docs/decisions.md 補 2026-04-29 ~ 2026-05-01 三筆 ADR
□ Day 6  PR: docs(refactor): 拆解 CLAUDE.md 與 docs 對齊現況""")

    # 四、長期維護機制
    d.add_page_break()
    add_heading(d, "四、長期維護機制(避免再爛掉)", 1)

    add_heading(d, "4.1 .gitignore 補強", 2)
    add_ascii_block(d, """\
# 在現有 .gitignore 末尾加上(避免下次 dev artefact 又進 repo)

# 一次性 LLM batch 產物
_retitle_proposals/
_resummary_proposals/
_*_proposals/

# 一次性 audit / review CSV/TXT
05_scripts/_*.csv
05_scripts/_*.txt
!05_scripts/_skip.txt   # 只有 _skip.txt 算「準輸入」,需 tracked

# 前端開發樣本
04_web/data/*_sample.json
04_web/data/_*.json     # 以 _ 開頭的暫存""")

    add_heading(d, "4.2 命名慣例(已實行,寫下來鞏固)", 2)
    add_bullet(d, "05_scripts/_*.py = 一次性整理腳本(預設 dry-run + --apply)")
    add_bullet(d, "00_source/_*.csv|.txt = 準輸入配置檔(_manifest / _skip)")
    add_bullet(d, "docs/_*.md = 一次性 handoff / 內部報告(可封存)")
    add_bullet(d, "docs/_*.docx = 規劃文件(對外公開度低)")
    add_bullet(d, "04_web/data/_*.json = 暫存(不可進 git)")

    add_heading(d, "4.3 季度檢查 SOP", 2)
    add_para(d, "每季跑一次本報告的「現況統計」,確認:")
    add_bullet(d, "untracked 區是否堆積 dev artefact")
    add_bullet(d, "tracked 是否有 N 個版本檔同時存在")
    add_bullet(d, "CLAUDE.md 是否又超過 30K(超過就拆)")
    add_bullet(d, "docs/ 內是否有 6 個月以上未動的 handoff(封存到 _archive/)")
    add_bullet(d, "_meta.json 的 node_count 是否與 README 同步")

    add_heading(d, "4.4 自動化建議(可選)", 2)
    add_bullet(d, "GitHub Action: stale-handoff-check — 每月跑,docs/_handoff_*.md 若 6 個月未動,開 issue 提醒封存")
    add_bullet(d, "GitHub Action: claudemd-size-check — CLAUDE.md > 50 KB 時 fail CI(因 gitignored 仍可在 PR 中提醒人工)")
    add_bullet(d, "Pre-commit hook: 拒絕 untracked 開頭為 _ 的檔進 staging")

    # 結語
    add_heading(d, "結語", 1)
    add_para(d, "資料夾本身的整潔度,是一個專案紀律的縮影。本資料庫已建立完整的腳本與文件體系,但累積一年多的迭代使得「過渡期保留」的心智成本日益增加。")
    add_para(d, "本報告建議分三批分週執行,每批可獨立回滾、不影響線上版,理想情況下三週後可達到:")
    add_bullet(d, "📦 本機釋出 ~2 MB,git 釋出 ~600 KB", color="2E7D4F")
    add_bullet(d, "📁 04_web/ 從 4 份 index.html 簡化為 1 份", color="2E7D4F")
    add_bullet(d, "📄 CLAUDE.md 從 92 KB → 25-30 KB", color="2E7D4F")
    add_bullet(d, "📚 docs/ 與現況對齊,_handoff 封存到 _archive/", color="2E7D4F")
    add_bullet(d, "🔧 .gitignore 強化,慣例落實到 hook,長期可自我維持", color="2E7D4F")
    add_para(d, "─── 報告完 ───", align=WD_ALIGN_PARAGRAPH.CENTER, color="888888", italic=True)

    d.save(out_path)
    print(f"  ✓ {out_path.relative_to(ROOT)}")


# ─────────────────────────────────────────────────────────────
# main
# ─────────────────────────────────────────────────────────────

def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    print("產出兩份報告...")
    build_flow_report(OUT_DIR / "_資料庫處理流程優化報告_2026-05-01.docx")
    build_cleanup_report(OUT_DIR / "_資料庫資料夾整理報告_2026-05-01.docx")
    print("完成。")


if __name__ == "__main__":
    main()
